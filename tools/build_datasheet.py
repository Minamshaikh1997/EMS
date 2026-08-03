from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "EMS_Product_Data_Sheet.docx"
NAVY, BLUE, PALE, WHITE, GRAY, GREEN = "17365D", "2F75B5", "F3F6FA", "FFFFFF", "666666", "1F7A5A"


def font(run, size=10, bold=False, color="111111", italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size, run.bold, run.italic = Pt(size), bold, italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    shd = pr.find(qn("w:shd")) or OxmlElement("w:shd")
    if shd.getparent() is None:
        pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, v=100, h=120):
    pr = cell._tc.get_or_add_tcPr()
    mar = pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar"); pr.append(mar)
    for name, val in (("top", v), ("start", h), ("bottom", v), ("end", h)):
        el = mar.find(qn("w:" + name))
        if el is None:
            el = OxmlElement("w:" + name); mar.append(el)
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")


def widths(table, values):
    table.autofit = False
    pr = table._tbl.tblPr
    tw = pr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tw.getparent() is None: pr.append(tw)
    tw.set(qn("w:w"), str(int(sum(values)*1440))); tw.set(qn("w:type"), "dxa")
    ind = pr.find(qn("w:tblInd")) or OxmlElement("w:tblInd")
    if ind.getparent() is None: pr.append(ind)
    ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for x in list(grid): grid.remove(x)
    for value in values:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(int(value*1440))); grid.append(col)
    for row in table.rows:
        for cell, value in zip(row.cells, values):
            cell.width = Inches(value); margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            pr = cell._tc.get_or_add_tcPr(); cw = pr.find(qn("w:tcW")) or OxmlElement("w:tcW")
            if cw.getparent() is None: pr.append(cw)
            cw.set(qn("w:w"), str(int(value*1440))); cw.set(qn("w:type"), "dxa")


def table(doc, headers, rows, col_widths, fs=9.1):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers): t.rows[0].cells[i].text = h
    hdr_pr = t.rows[0]._tr.get_or_add_trPr(); rep = OxmlElement("w:tblHeader"); rep.set(qn("w:val"), "true"); hdr_pr.append(rep)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row): cells[i].text = value
    widths(t, col_widths)
    for ri, row in enumerate(t.rows):
        for cell in row.cells:
            shade(cell, NAVY if ri == 0 else (PALE if ri % 2 == 0 else WHITE))
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
                for r in p.runs: font(r, fs, ri == 0, WHITE if ri == 0 else "111111")
    return t


def para(doc, text, size=10.2, bold=False, color="111111", after=5, align=None, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.1
    if align is not None: p.alignment = align
    font(p.add_run(text), size, bold, color, italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}"); p.paragraph_format.keep_with_next = True; p.add_run(text); return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.05
    font(p.add_run(text), 9.8); return p


doc = Document(); sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = Inches(1)
sec.top_margin, sec.bottom_margin = Inches(.72), Inches(.68)
sec.header_distance, sec.footer_distance = Inches(.28), Inches(.28)

normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(10.2)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
for name, size, color, before, after in (("Heading 1", 15, NAVY, 10, 5), ("Heading 2", 11.5, BLUE, 7, 3)):
    st = doc.styles[name]; st.font.name = "Arial"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    st.paragraph_format.space_before, st.paragraph_format.space_after = Pt(before), Pt(after)

hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(hp.add_run("EMS  |  PRODUCT DATA SHEET  |  VERSION 2.0"), 7.8, True, GRAY)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("Employee Management System | Technical information subject to final deployment configuration"), 7.7, False, GRAY)

para(doc, "PRODUCT DATA SHEET", 10.5, True, BLUE, 8, WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "Employee Management System", 25, True, NAVY, 5, WD_ALIGN_PARAGRAPH.CENTER)
para(doc, "Web-based HR operations, attendance, leave, payroll and access control", 12.5, False, BLUE, 18, WD_ALIGN_PARAGRAPH.CENTER)

table(doc, ["Product", "Deployment", "Platform", "Database"], [["EMS v2.0", "On-premise / hosted web", "PHP web application", "MySQL / utf8mb4"]], [1.25, 2.1, 1.75, 1.4], 8.8)
para(doc, "", 1, after=5)
table(doc, ["Primary users", "Best suited for", "Commercial model"], [["HR, management, supervisors and employees", "Small and mid-sized organizations", "Custom deployment with source-code handover"]], [2.2, 2.2, 2.1], 8.8)

heading(doc, "Product Overview", 1)
para(doc, "EMS is a centralized browser-based system for maintaining employee records and operating day-to-day HR workflows. Separate admin and employee portals combine attendance, shifts, leave, payroll, reporting and granular permissions in one PHP/MySQL application.")

heading(doc, "Key Capabilities", 1)
capabilities = [
    ("Employee management", "Create, edit, activate/deactivate, search and maintain employee profiles and photographs."),
    ("Attendance", "Employee attendance marking, history, admin reporting and correction/adjustment requests."),
    ("Shift management", "Configure work shifts and retain shift-history support."),
    ("Leave management", "Leave application, balance, history, administrative review, approval and rejection."),
    ("Payroll", "Salary components, employee salary structure, monthly payroll generation, slips and history."),
    ("Reports & exports", "Attendance and payroll reporting, printable salary slips and Excel export."),
    ("Role-based access", "Eight-level role hierarchy with view, create, edit, delete, approve and export permissions."),
    ("Employee rights", "Per-employee controls for payroll, leave, attendance, adjustments, profile, reports and password access."),
    ("Communications", "Administration of notices and holidays; email utility available for configured workflows."),
]
table(doc, ["Module", "Function"], capabilities, [1.75, 4.75], 8.8)

doc.add_page_break()
heading(doc, "Functional Architecture", 1)
table(doc, ["Portal", "Users", "Core functions"], [
    ("Admin portal", "Super Admin, Admin, VP, Operations Manager, WFM, Supervisor, Team Lead", "Dashboards, employees, shifts, attendance, leave approvals, payroll, reports, rights and role permissions"),
    ("Employee portal", "Employee users", "Dashboard, attendance, leave, payroll visibility, adjustments, profile/photo and password"),
], [1.25, 2.2, 3.05], 8.7)

heading(doc, "Roles & Permission Model", 1)
table(doc, ["Role", "Hierarchy", "Typical responsibility"], [
    ("Super Admin", "1", "Unrestricted system administration"),
    ("Admin", "2", "Organization-level administration"),
    ("VP / Operations Manager / WFM", "3", "Operational management and approved module access"),
    ("Supervisor", "4", "Team supervision and workflow review"),
    ("Team Lead", "5", "Limited team operations"),
    ("Employee", "6", "Personal self-service functions"),
], [2.25, 1.0, 3.25], 8.8)
para(doc, "Permissions are configurable per module across View, Create, Edit, Delete, Approve and Export actions. Super Admin has an unrestricted bypass. Employee-level rights provide an additional layer for seven self-service features.", 9.5, False, GRAY, 5, italic=True)

heading(doc, "Technical Specifications", 1)
table(doc, ["Item", "Specification"], [
    ("Application type", "Responsive browser-based web application"),
    ("Backend", "Core PHP with MySQLi database access"),
    ("Database", "MySQL; utf8mb4 connection character set"),
    ("Frontend", "HTML5, CSS, JavaScript and Bootstrap-style responsive UI"),
    ("Server timezone", "Asia/Karachi"),
    ("Local environment", "Apache + PHP + MySQL; XAMPP-compatible"),
    ("Authentication", "Separate admin and employee login/session flows"),
    ("Exports", "Excel-compatible export and printable salary slips"),
    ("Storage", "Relational HR/payroll data plus employee image uploads"),
], [1.75, 4.75], 8.9)

heading(doc, "Database Coverage", 1)
for item in [
    "Employee, admin and department records",
    "Attendance, shifts and shift-history records",
    "Leave requests and leave balances",
    "Attendance adjustment requests",
    "Salary components, structures, payroll runs and salary slips",
    "Roles, permissions, role mappings and permission audit logs",
    "Notices, holidays and employee feature-right fields",
]: bullet(doc, item)

doc.add_page_break()
heading(doc, "Security & Control Features", 1)
for item in [
    "Prepared statements are used in core permission and workflow areas to reduce SQL-injection risk.",
    "CSRF token handling is implemented in the role/permission module.",
    "Server-side role and permission checks support restricted routes and conditional actions.",
    "Permission activity can be recorded in an audit-log table.",
    "Password hashing and verification utilities are present for credential migration and authentication.",
    "Production hardening is required before public launch: remove embedded credentials, restrict setup/reset/test routes, validate uploads and complete a route-by-route authorization review.",
]: bullet(doc, item)

heading(doc, "Recommended System Requirements", 1)
table(doc, ["Component", "Minimum", "Recommended production"], [
    ("Web server", "Apache 2.4", "Current supported Apache or Nginx"),
    ("PHP", "PHP 8.0+", "PHP 8.2/8.3 with MySQLi"),
    ("Database", "MySQL 5.7+", "MySQL 8.0+ with automated backups"),
    ("Memory", "2 GB RAM", "4 GB+ RAM for typical SMB use"),
    ("Storage", "10 GB", "SSD storage sized for uploads and backups"),
    ("Client browser", "Modern Chrome/Edge/Firefox", "Latest stable desktop/mobile browser"),
    ("Network", "LAN or internet", "HTTPS with stable connectivity"),
], [1.35, 2.1, 3.05], 8.7)

heading(doc, "Deployment & Integration", 1)
for item in [
    "Supports local network, VPS or compatible shared-hosting deployment subject to server limits.",
    "Requires a configured MySQL database, writable uploads directory and secure environment-specific credentials.",
    "SMTP/email, biometric devices, SMS/WhatsApp and third-party systems require separate configuration or integration work.",
    "Database migration/setup scripts are included but should be consolidated into a controlled production release process.",
]: bullet(doc, item)

heading(doc, "Current Scope Boundaries", 1)
table(doc, ["Included in current web product", "Requires extension / validation"], [
    ("Employee, attendance, shifts, leave, payroll, reporting, RBAC and employee rights", "Mobile applications, recruitment/ATS, performance reviews, asset management and multi-company tenancy"),
    ("Salary structures, components and payroll slips", "Verified FBR tax, EOBI, PESSI/SESSI and province-specific statutory calculations"),
    ("Web attendance and adjustment workflow", "Biometric hardware, geofencing, face recognition and offline attendance"),
], [3.25, 3.25], 8.5)

heading(doc, "Product Status", 1)
para(doc, "Status: Functional development codebase requiring final consolidation, security hardening, regression testing, controlled deployment and client acceptance before production use.", 10.2, True, GREEN, 5)
para(doc, "Document date: 03 August 2026 | Data sheet version: 1.0 | Product version referenced: EMS v2.0+", 8.8, False, GRAY, 0, WD_ALIGN_PARAGRAPH.CENTER)

doc.core_properties.title = "EMS Product Data Sheet"
doc.core_properties.subject = "Technical and functional data sheet for Employee Management System"
doc.core_properties.author = "Software Development Team"
doc.core_properties.keywords = "EMS, HRMS, data sheet, PHP, MySQL"
doc.save(OUT)
print(OUT)
