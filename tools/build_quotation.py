from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "EMS_Project_Quotation_Pakistan.docx"

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "DCE6F1"
PALE = "F3F6FA"
GOLD = "D6A84B"
GREEN = "1F7A5A"
GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "111111"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_inches):
    table.autofit = False
    total = sum(widths_inches)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(total * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, (cell, w) in enumerate(zip(row.cells, widths_inches)):
            cell.width = Inches(w)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(w * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, color=None, italic=None, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", size=10.5, bold=False, color=BLACK, align=None,
             before=0, after=6, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_font(r, size=10.2, color=BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [6.5])
    c = table.cell(0, 0)
    set_cell_shading(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_font(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_font(r, size=10.5, color=BLACK)
    add_para(doc, "", size=1, after=5)


def format_table(table, header=True, font_size=9.3):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(table.rows):
        if ri == 0 and header:
            set_repeat_table_header(row)
        for cell in row.cells:
            if ri == 0 and header:
                set_cell_shading(cell, NAVY)
            elif ri % 2 == 0:
                set_cell_shading(cell, PALE)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    set_font(r, size=font_size, bold=(ri == 0 and header), color=WHITE if (ri == 0 and header) else BLACK)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.3)
sec.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12
for name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 12, 6),
    ("Heading 2", 12.5, BLUE, 9, 4),
    ("Heading 3", 10.8, NAVY, 6, 2),
]:
    st = styles[name]
    st.font.name = "Arial"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# Running header/footer
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("EMPLOYEE MANAGEMENT SYSTEM  |  COMMERCIAL QUOTATION")
set_font(hr, size=7.8, bold=True, color=GRAY)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Confidential | Valid for 15 days from 03 August 2026")
set_font(fr, size=8, color=GRAY)

# First-page proposal centerpiece
add_para(doc, "COMMERCIAL PROPOSAL", size=10.5, bold=True, color=GOLD,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=20, after=10)
add_para(doc, "Employee Management System (EMS)", size=25, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
add_para(doc, "Completion, Production Hardening, Deployment & Handover",
         size=13.5, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)

meta = doc.add_table(rows=4, cols=2)
set_table_widths(meta, [3.25, 3.25])
meta_data = [
    ("Prepared for", "Client / Organization"),
    ("Prepared by", "Software Development Team"),
    ("Quotation date", "03 August 2026"),
    ("Currency", "Pakistani Rupees (PKR)"),
]
for row, (a, b) in zip(meta.rows, meta_data):
    row.cells[0].text = a
    row.cells[1].text = b
for ri, row in enumerate(meta.rows):
    for ci, cell in enumerate(row.cells):
        set_cell_shading(cell, PALE if ri % 2 == 0 else WHITE)
        for r in cell.paragraphs[0].runs:
            set_font(r, size=10, bold=(ci == 0), color=NAVY if ci == 0 else BLACK)
format_table(meta, header=False, font_size=10)

add_para(doc, "", size=1, after=13)
add_callout(doc, "RECOMMENDED FIXED PRICE", "PKR 425,000 (Four Hundred Twenty-Five Thousand Rupees)", fill="E8F3EE", accent=GREEN)
add_para(doc, "Estimated delivery: 7-8 weeks after kickoff, access handover and scope confirmation.",
         size=10.5, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
add_para(doc, "This quotation is based on a technical review of the supplied PHP/MySQL codebase and Pakistan's local custom-development market. It prices the work needed to convert the current feature-rich build into a stable, testable and deployable client delivery.",
         size=10.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10)

doc.add_page_break()
add_heading(doc, "1. Executive Summary", 1)
add_para(doc, "The reviewed project is a browser-based Employee Management System built in core PHP and MySQL. It already contains substantial functional code across two portals. The recommended engagement therefore focuses on consolidation, defect correction, security hardening, workflow validation, production deployment and client handover rather than rebuilding every module from zero.")
add_callout(doc, "COMMERCIAL POSITION", "A fair local-market quote for the reviewed scope is PKR 425,000. A typical clean-sheet custom build of the same breadth would normally be priced higher because discovery, architecture and all modules would start from zero.")

add_heading(doc, "2. Reviewed Project Scope", 1)
scope = [
    ("Admin portal", "Dashboard, employees, departments, notices, holidays and user administration"),
    ("Employee self-service", "Employee login, dashboard, profile, photo, password and personal records"),
    ("Attendance & shifts", "Attendance marking/history, shift management, reports and adjustment requests"),
    ("Leave management", "Apply, balance, history, admin review, approval and rejection workflows"),
    ("Payroll", "Salary structure/components, monthly generation, slips, history and payroll reports"),
    ("Access control", "Eight-role hierarchy, module-level RBAC, audit logs and per-employee feature rights"),
    ("Reporting", "Attendance/payroll reporting, print views and Excel export"),
    ("Database utilities", "Setup, migrations, verification, diagnostic and repair scripts"),
]
t = doc.add_table(rows=1, cols=2)
t.rows[0].cells[0].text = "Area"
t.rows[0].cells[1].text = "Included capability"
for a, b in scope:
    cells = t.add_row().cells
    cells[0].text = a
    cells[1].text = b
set_table_widths(t, [1.65, 4.85])
format_table(t)

add_heading(doc, "3. Technical Review Snapshot", 1)
for item in [
    "Core stack: PHP, MySQL, HTML/CSS/JavaScript, Bootstrap-style responsive interfaces and XAMPP-oriented local setup.",
    "Code footprint: more than 100 project files, with approximately 18,000+ lines across PHP, SQL, CSS and documentation.",
    "Automated syntax check: all PHP files passed linting during review.",
    "The repository contains production credentials/placeholders and browser-accessible setup, reset, diagnostic and test utilities; these must be removed or protected before launch.",
    "Several zero-length legacy/alias files and overlapping old/current pages indicate a consolidation and regression-testing pass is required.",
    "The delivered scope does not yet demonstrate automated tests, CI/CD, backup automation, privacy controls or verified Pakistani statutory payroll calculations.",
]:
    add_bullet(doc, item)

doc.add_page_break()
add_heading(doc, "4. Deliverables", 1)
deliverables = [
    ("Codebase stabilization", "Remove dead/duplicate routes, reconcile setup scripts, fix confirmed defects and standardize configuration."),
    ("Security hardening", "Protect credentials, enforce server-side authorization, secure uploads, review sessions/CSRF and disable public diagnostic/reset endpoints."),
    ("Workflow completion", "Validate employee, attendance, leave, adjustment, payroll and role/permission workflows end to end."),
    ("UI/UX polish", "Correct responsive issues, navigation consistency, validation messages, print layouts and dark-mode behavior where applicable."),
    ("Database readiness", "Produce a clean schema/migration path, indexes, seed guidance and a backup/restore checklist."),
    ("Quality assurance", "Functional and role-based testing on current desktop browsers plus representative mobile widths."),
    ("Deployment", "One production deployment on client-provided hosting/VPS with SSL and database configuration."),
    ("Handover", "Admin walkthrough, deployment notes, known-limitations register and 30 days of post-launch defect support."),
]
for title, desc in deliverables:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title + ": ")
    set_font(r, size=10.4, bold=True, color=NAVY)
    r = p.add_run(desc)
    set_font(r, size=10.4, color=BLACK)

add_heading(doc, "5. Commercial Breakdown", 1)
items = [
    ("1", "Technical audit, scope freeze & stabilization plan", "35,000"),
    ("2", "Core configuration, database cleanup & migrations", "45,000"),
    ("3", "Employee, attendance, shifts & adjustment workflows", "75,000"),
    ("4", "Leave management and approval workflows", "40,000"),
    ("5", "Payroll, salary structure, slips & reports validation", "80,000"),
    ("6", "RBAC, employee rights & security hardening", "60,000"),
    ("7", "UI polish, regression QA & acceptance fixes", "50,000"),
    ("8", "Deployment, documentation, training & 30-day warranty", "40,000"),
]
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(("#", "Work package", "Amount (PKR)")):
    t.rows[0].cells[i].text = h
for a, b, c in items:
    cells = t.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = a, b, c
cells = t.add_row().cells
cells[0].merge(cells[1]).text = "TOTAL FIXED PROJECT FEE"
cells[2].text = "425,000"
set_table_widths(t, [0.45, 4.65, 1.4])
format_table(t)
for r in t.rows[-1].cells:
    set_cell_shading(r, "E8F3EE")
    for run in r.paragraphs[0].runs:
        set_font(run, size=10.2, bold=True, color=GREEN)

add_para(doc, "Taxes: Any applicable withholding or sales tax will be handled according to the client's legal status and prevailing Pakistani law. The quoted project fee is exclusive of third-party charges.", size=9.3, italic=True, color=GRAY, before=5)

doc.add_page_break()
add_heading(doc, "6. Delivery Plan", 1)
timeline = [
    ("Week 1", "Kickoff, environment setup, scope freeze and issue register"),
    ("Weeks 2-3", "Database/configuration cleanup, employee, attendance and leave workflows"),
    ("Weeks 4-5", "Payroll validation, RBAC/rights enforcement and security hardening"),
    ("Weeks 6-7", "UI polish, regression QA, client UAT and acceptance fixes"),
    ("Week 8", "Production deployment, training and handover buffer"),
]
t = doc.add_table(rows=1, cols=2)
t.rows[0].cells[0].text = "Period"
t.rows[0].cells[1].text = "Milestone"
for a, b in timeline:
    cells = t.add_row().cells
    cells[0].text, cells[1].text = a, b
set_table_widths(t, [1.25, 5.25])
format_table(t)

add_heading(doc, "7. Payment Schedule", 1)
payments = [
    ("40%", "PKR 170,000", "Advance on approval and kickoff"),
    ("30%", "PKR 127,500", "After core workflows are ready for client UAT"),
    ("20%", "PKR 85,000", "On UAT completion and production deployment"),
    ("10%", "PKR 42,500", "On final handover; due within 7 days"),
]
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(("Stage", "Amount", "Trigger")):
    t.rows[0].cells[i].text = h
for a, b, c in payments:
    cells = t.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = a, b, c
set_table_widths(t, [0.9, 1.45, 4.15])
format_table(t)

add_heading(doc, "8. Assumptions & Acceptance", 1)
for item in [
    "Client will provide hosting/VPS, domain/SSL access, database access, email SMTP details, logo/brand assets and one authorized decision-maker.",
    "Scope covers one company/tenant and the workflows presently represented in the reviewed repository.",
    "Client UAT feedback will be consolidated and returned within three business days per review round.",
    "Two consolidated UAT correction rounds are included. New features or policy changes will be estimated separately through change control.",
    "Acceptance occurs when agreed workflows pass the signed checklist; cosmetic preferences outside the agreed design are change requests.",
]:
    add_bullet(doc, item)

doc.add_page_break()
add_heading(doc, "9. Exclusions & Optional Add-ons", 1)
excluded = [
    ("Biometric device integration", "From 75,000", "Depends on device/API and onsite access"),
    ("Pakistani tax/EOBI/PESSI statutory engine", "From 60,000", "Requires policy confirmation and accountant validation"),
    ("Data migration from Excel/legacy software", "From 25,000", "Based on data volume and cleanliness"),
    ("Mobile apps (Android/iOS)", "Separate estimate", "Not included in current web scope"),
    ("Cloud hosting, domain, SSL, email/SMS/WhatsApp", "At actual cost", "Third-party recurring charges"),
    ("Ongoing maintenance", "PKR 25,000/month", "Up to 12 support hours; new modules excluded"),
]
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(("Item", "Indicative price", "Note")):
    t.rows[0].cells[i].text = h
for a, b, c in excluded:
    cells = t.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = a, b, c
set_table_widths(t, [2.35, 1.35, 2.8])
format_table(t, font_size=8.9)

add_heading(doc, "10. Warranty, Ownership & Validity", 1)
for item in [
    "Warranty: 30 calendar days after production launch for reproducible defects within the agreed scope.",
    "Ownership: Source code and project-specific deliverables transfer after full payment. Third-party libraries remain under their original licenses.",
    "Confidentiality: Client data and credentials will be treated as confidential and used only for delivery/support.",
    "Validity: This quotation remains valid for 15 days from 03 August 2026.",
    "Schedule: Timeline starts after advance payment, access handover and written scope approval.",
]:
    add_bullet(doc, item)

add_heading(doc, "11. Market Basis", 1)
add_para(doc, "The price has been calibrated against current Pakistan-market HRMS subscription benchmarks and local development economics, then adjusted for the fact that a substantial codebase already exists. Public 2026 benchmarks commonly place full HRMS subscriptions around PKR 300-1,000 per employee per month, while local products advertise plans ranging from roughly PKR 3,000 to PKR 120,000 per month depending on headcount and depth. A custom, source-code-owned implementation is priced differently because it includes one-off engineering, deployment risk and handover.")
add_para(doc, "Reference points reviewed: Indeed Pakistan web-developer salary data (updated July 2026); eHR Pakistan pricing; QuoHR pricing; MAB HRMS Pakistan pricing; and published Pakistan HRMS pricing guides. These are market indicators, not direct substitutes for this custom fixed-price scope.", size=9.2, italic=True, color=GRAY)

add_heading(doc, "Approval", 1)
approval = doc.add_table(rows=3, cols=2)
set_table_widths(approval, [3.25, 3.25])
labels = [("Client name / company", ""), ("Authorized signature", ""), ("Date", "")]
for row, (a, b) in zip(approval.rows, labels):
    row.cells[0].text = a
    row.cells[1].text = b
    for cell in row.cells:
        set_cell_shading(cell, PALE)
        set_cell_margins(cell, top=180, bottom=180)
    for r in row.cells[0].paragraphs[0].runs:
        set_font(r, size=9.5, bold=True, color=NAVY)

# Core properties
doc.core_properties.title = "Employee Management System - Commercial Quotation"
doc.core_properties.subject = "Pakistan local-market quotation for EMS completion and deployment"
doc.core_properties.author = "Software Development Team"
doc.core_properties.keywords = "EMS, HRMS, quotation, Pakistan, PHP, MySQL"

doc.save(OUT)
print(OUT)
